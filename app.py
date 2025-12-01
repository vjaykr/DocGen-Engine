

import os
import re
import json
import threading
import traceback
from datetime import datetime
from pathlib import Path
import pandas as pd
from docxtpl import DocxTemplate
from docx import Document
import win32com.client
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from tkinter import font as tkFont
import ctypes
from ctypes import wintypes

class SaralWorksApp:
    def __init__(self):
        self.root = tk.Tk()
        self.setup_window()
        self.setup_variables()
        self.setup_styles()
        
        # Backend data
        self.template_path = ""
        self.excel_path = ""
        self.output_folder = str(Path.home() / "Desktop")
        self.placeholders = []
        self.columns = []
        self.mapping = {}
        self.generate_excel_btn = None
        
        self.create_widgets()
        
    def setup_window(self):
        # Enable DPI awareness for sharp fonts
        try:
            ctypes.windll.shcore.SetProcessDpiAwareness(1)
        except:
            pass
            
        self.root.title("ENGINEER")
        self.root.geometry("1400x800")
        self.root.configure(bg='#f8fafc')
        self.root.resizable(True, True)
        self.root.state('zoomed')  # Maximize window on startup
        
        # Configure for high DPI
        self.root.tk.call('tk', 'scaling', 1.5)
        
    def setup_variables(self):
        self.status_var = tk.StringVar(value="Ready to start")
        self.template_name_var = tk.StringVar(value="")
        self.excel_name_var = tk.StringVar(value="")
        self.output_folder_var = tk.StringVar(value="Desktop")
        self.include_mobile_var = tk.BooleanVar(value=False)
        self.generating_var = tk.BooleanVar(value=False)
        
    def setup_styles(self):
        self.style = ttk.Style()
        self.style.theme_use('clam')
        
        # Configure styles for glass morphism
        self.style.configure('Glass.TButton', 
                           background='rgba(255,255,255,0.2)',
                           borderwidth=1,
                           relief='flat',
                           font=('Segoe UI', 10))
        
        # Configure combobox style
        self.style.configure('Glass.TCombobox',
                           fieldbackground='white',
                           borderwidth=1,
                           relief='solid')
        
    def create_widgets(self):
        # Main container
        main_frame = ttk.Frame(self.root)
        main_frame.pack(fill='both', expand=True, padx=10, pady=10)
        
        # Header
        self.create_header(main_frame)
        
        # Content area
        content_frame = ttk.Frame(main_frame)
        content_frame.pack(fill='both', expand=True, pady=(10, 0))
        
        # Left sidebar
        self.create_sidebar(content_frame)
        
        # Right content area
        self.create_main_content(content_frame)
        
        # Bottom generate button
        self.create_generate_button(main_frame)
        
    def create_header(self, parent):
        header_frame = tk.Frame(parent, bg='#1e40af', height=80)
        header_frame.pack(fill='x', pady=(0, 10))
        header_frame.pack_propagate(False)
        
        # Title (H1: 22px)
        title_label = tk.Label(header_frame, text="DocGen Engine", 
                              font=('Segoe UI', 20, 'bold'), 
                              fg='white', bg='#1e40af')
        title_label.pack(side='left', padx=20, pady=20)
        
        # # Subtitle (Caption: 11px)
        # subtitle_label = tk.Label(header_frame, text="Contract AutoFiller & PDF Generator", 
        #                          font=('Segoe UI', 9), 
        #                          fg='#e0e7ff', bg='#1e40af')
        # subtitle_label.pack(side='left', padx=(0, 20), pady=(28, 20))
        
        # Status (Body: 12px)
        status_label = tk.Label(header_frame, textvariable=self.status_var,
                               font=('Segoe UI', 12, 'bold'), 
                               fg='white', bg='#1e40af')
        status_label.pack(side='right', padx=20, pady=20)
        
    def create_sidebar(self, parent):
        sidebar_frame = tk.Frame(parent, bg='#f8fafc', width=380)
        sidebar_frame.pack(side='left', fill='y', padx=(0, 10))
        sidebar_frame.pack_propagate(False)
        
        # File Selection Card
        self.create_file_selection_card(sidebar_frame)
        
        # Actions Card
        self.create_actions_card(sidebar_frame)
        
        # Settings Card
        self.create_settings_card(sidebar_frame)
        
    def create_file_selection_card(self, parent):
        # Glass morphism card with rounded corners effect
        card_frame = tk.Frame(parent, bg='white', relief='flat', bd=0)
        card_frame.pack(fill='x', padx=15, pady=10)
        
        # Add shadow effect with nested frame
        shadow_frame = tk.Frame(card_frame, bg='#e5e7eb', height=2)
        shadow_frame.pack(fill='x', side='bottom')
        
        # Main card content
        content_frame = tk.Frame(card_frame, bg='white')
        content_frame.pack(fill='both', expand=True, padx=2, pady=2)
        
        # Card header (H2: 18px)
        header_label = tk.Label(content_frame, text="📁 File Selection", 
                               font=('Segoe UI', 18, 'bold'), 
                               bg='white', fg='#1f2937')
        header_label.pack(anchor='w', padx=15, pady=(15, 12))
        
        # Template selection (Body: 14px)
        tk.Label(content_frame, text="Word Template", font=('Segoe UI', 14, 'bold'), 
                bg='white', fg='#374151').pack(anchor='w', padx=15, pady=(10, 5))
        
        template_frame = tk.Frame(content_frame, bg='white')
        template_frame.pack(fill='x', padx=15, pady=(0, 10))
        
        self.template_entry = tk.Entry(template_frame, textvariable=self.template_name_var, 
                                      state='readonly', font=('Segoe UI', 14), 
                                      relief='solid', bd=1)
        self.template_entry.pack(side='left', fill='x', expand=True, padx=(0, 8), ipady=6)
        
        # Glass morphism button (Body: 16px)
        template_btn = tk.Button(template_frame, text="📄", font=('Segoe UI', 16), 
                                bg='#3b82f6', fg='white', width=3, height=1,
                                relief='flat', bd=0, cursor='hand2',
                                command=self.select_template)
        template_btn.pack(side='right')
        
        # Excel selection (Body: 14px)
        tk.Label(content_frame, text="Excel Data", font=('Segoe UI', 14, 'bold'), 
                bg='white', fg='#374151').pack(anchor='w', padx=15, pady=(10, 5))
        
        excel_frame = tk.Frame(content_frame, bg='white')
        excel_frame.pack(fill='x', padx=15, pady=(0, 10))
        
        self.excel_entry = tk.Entry(excel_frame, textvariable=self.excel_name_var, 
                                   state='readonly', font=('Segoe UI', 14),
                                   relief='solid', bd=1)
        self.excel_entry.pack(side='left', fill='x', expand=True, padx=(0, 8), ipady=6)
        
        excel_btn = tk.Button(excel_frame, text="📊", font=('Segoe UI', 16), 
                             bg='#10b981', fg='white', width=3, height=1,
                             relief='flat', bd=0, cursor='hand2',
                             command=self.select_excel)
        excel_btn.pack(side='right')
        
        # Output folder selection (Body: 14px)
        tk.Label(content_frame, text="Output Folder", font=('Segoe UI', 14, 'bold'), 
                bg='white', fg='#374151').pack(anchor='w', padx=15, pady=(10, 5))
        
        folder_frame = tk.Frame(content_frame, bg='white')
        folder_frame.pack(fill='x', padx=15, pady=(0, 10))
        
        self.folder_entry = tk.Entry(folder_frame, textvariable=self.output_folder_var, 
                                    state='readonly', font=('Segoe UI', 14),
                                    relief='solid', bd=1)
        self.folder_entry.pack(side='left', fill='x', expand=True, padx=(0, 8), ipady=6)
        
        folder_btn = tk.Button(folder_frame, text="📁", font=('Segoe UI', 16), 
                              bg='#f59e0b', fg='white', width=3, height=1,
                              relief='flat', bd=0, cursor='hand2',
                              command=self.select_folder)
        folder_btn.pack(side='right')
        
        # Scan button (Body: 15px)
        self.scan_btn = tk.Button(content_frame, text="🔍 Scan Placeholders", 
                                 font=('Segoe UI', 15, 'bold'), 
                                 bg='#7c3aed', fg='white', height=2,
                                 relief='flat', bd=0, cursor='hand2',
                                 command=self.scan_placeholders)
        self.scan_btn.pack(fill='x', padx=15, pady=(12, 15))
        
    def create_actions_card(self, parent):
        # Glass morphism card
        card_frame = tk.Frame(parent, bg='white', relief='flat', bd=0)
        card_frame.pack(fill='x', padx=15, pady=10)
        
        shadow_frame = tk.Frame(card_frame, bg='#e5e7eb', height=2)
        shadow_frame.pack(fill='x', side='bottom')
        
        content_frame = tk.Frame(card_frame, bg='white')
        content_frame.pack(fill='both', expand=True, padx=2, pady=2)
        
        # Card header (H2: 18px)
        header_label = tk.Label(content_frame, text="⚙️ Actions", 
                               font=('Segoe UI', 18, 'bold'), 
                               bg='white', fg='#1f2937')
        header_label.pack(anchor='w', padx=15, pady=(15, 12))
        
        action_grid = tk.Frame(content_frame, bg='white')
        action_grid.pack(fill='x', padx=15, pady=(0, 12))
        for col in range(2):
            action_grid.columnconfigure(col, weight=1)
        
        # Auto map button (Body: 14px)
        self.auto_map_btn = tk.Button(action_grid, text="Auto Map Fields", 
                         font=('Segoe UI', 11, 'bold'), 
                                     bg='#10b981', fg='white',
                                     relief='flat', bd=0, cursor='hand2',
                                     command=self.auto_map)
        self.auto_map_btn.grid(row=0, column=0, padx=(0, 8), pady=0, sticky='nsew', ipady=12)
        
        # Clear map button (Body: 14px
        self.clear_btn = tk.Button(action_grid, text="Clear Mapping", 
                      font=('Segoe UI', 11, 'bold'), 
                                  bg='#ef4444', fg='white',
                                  relief='flat', bd=0, cursor='hand2',
                                  command=self.clear_mapping)
        self.clear_btn.grid(row=0, column=1, padx=(8, 0), pady=0, sticky='nsew', ipady=12)
        
        # Preview button (Body: 14px)
        self.preview_btn = tk.Button(content_frame, text="👁️ Preview Data", 
                        font=('Segoe UI', 11, 'bold'), 
                                    bg='#3b82f6', fg='white',
                                    relief='flat', bd=0, cursor='hand2',
                                    command=self.preview_data)
        self.preview_btn.pack(fill='x', padx=15, pady=(0, 15), ipady=10)
        
    def create_settings_card(self, parent):
        # Glass morphism card
        card_frame = tk.Frame(parent, bg='white', relief='flat', bd=0)
        card_frame.pack(fill='x', padx=15, pady=10)
        
        shadow_frame = tk.Frame(card_frame, bg='#e5e7eb', height=2)
        shadow_frame.pack(fill='x', side='bottom')
        
        content_frame = tk.Frame(card_frame, bg='white')
        content_frame.pack(fill='both', expand=True, padx=2, pady=2)
        
        # Card header (H2: 18px)
        header_label = tk.Label(content_frame, text="⚙️ Settings", 
                               font=('Segoe UI', 18, 'bold'), 
                               bg='white', fg='#1f2937')
        header_label.pack(anchor='w', padx=15, pady=(15, 12))
        
        # Mobile checkbox (Body: 14px)
        mobile_check = tk.Checkbutton(content_frame, text="Include Mobile in Filename", 
                                     variable=self.include_mobile_var,
                                     font=('Segoe UI', 14), bg='white', fg='#374151')
        mobile_check.pack(anchor='w', padx=15, pady=(0, 15))
        
    def create_main_content(self, parent):
        main_content_frame = tk.Frame(parent, bg='white', relief='solid', bd=1)
        main_content_frame.pack(side='right', fill='both', expand=True)
        
        # Header
        content_header = tk.Frame(main_content_frame, bg='#3b82f6', height=70)
        content_header.pack(fill='x')
        content_header.pack_propagate(False)
        
        header_label = tk.Label(content_header, text="🔗 Field Mapping", 
                               font=('Segoe UI', 17, 'bold'), 
                               fg='white', bg='#3b82f6')
        header_label.pack(side='left', padx=20, pady=15)
        
        self.field_count_label = tk.Label(content_header, text="0 fields → 0 columns", 
                                         font=('Segoe UI', 13), 
                                         fg='#e0e7ff', bg='#3b82f6')
        self.field_count_label.pack(side='right', padx=20, pady=15)
        
        # Mapping area
        self.create_mapping_area(main_content_frame)
        
    def create_mapping_area(self, parent):
        # Scrollable frame for mappings
        canvas = tk.Canvas(parent, bg='#f9fafb')
        scrollbar = ttk.Scrollbar(parent, orient="vertical", command=canvas.yview)
        self.scrollable_frame = tk.Frame(canvas, bg='#f9fafb')
        
        self.scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas.create_window((0, 0), window=self.scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        # Bind mouse wheel to canvas
        def on_mousewheel(event):
            canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        
        canvas.bind("<MouseWheel>", on_mousewheel)
        
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        # Initial empty state
        self.show_empty_state()
        
    def show_empty_state(self):
        for widget in self.scrollable_frame.winfo_children():
            widget.destroy()
            
        empty_frame = tk.Frame(self.scrollable_frame, bg='#f9fafb')
        empty_frame.pack(expand=True, fill='both')
        
        center_frame = tk.Frame(empty_frame, bg='#f9fafb')
        center_frame.place(relx=0.5, rely=0.5, anchor='center')
        
        tk.Label(center_frame, text="🔍", font=('Segoe UI', 48), 
                bg='#f9fafb', fg='#9ca3af').pack()
        tk.Label(center_frame, text="No Fields Detected", 
                font=('Segoe UI', 18, 'bold'), 
                bg='#f9fafb', fg='#9ca3af').pack(pady=(12, 6))
        tk.Label(center_frame, text="Select files and click 'Scan Placeholders' to begin", 
                font=('Segoe UI', 13), 
                bg='#f9fafb', fg='#6b7280').pack()
        
    def create_generate_button(self, parent):
        generate_frame = tk.Frame(parent, bg='#f8fafc', height=110)
        generate_frame.pack(fill='x', pady=(10, 0))
        generate_frame.pack_propagate(False)
        
        button_row = tk.Frame(generate_frame, bg='#f8fafc')
        button_row.pack(fill='x', padx=60, pady=20)
        button_row.columnconfigure(0, weight=3)
        button_row.columnconfigure(1, weight=2)
        
        self.generate_btn = tk.Button(button_row, text="🚀 Generate DOCX & PDF Files", 
                                     font=('Segoe UI', 16, 'bold'), 
                                     bg='#10b981', fg='white', height=2,
                                     relief='flat', bd=0, cursor='hand2',
                                     command=self.generate_documents)
        self.generate_btn.grid(row=0, column=0, sticky='nsew', padx=(0, 12))
        
        self.generate_excel_btn = tk.Button(button_row, text="📑 Generate Excel Report", 
                                           font=('Segoe UI', 14, 'bold'), 
                                           bg='#10b981', fg='white', height=2,
                                           relief='flat', bd=0, cursor='hand2',
                                           state='disabled', command=self.generate_excel_report)
        self.generate_excel_btn.grid(row=0, column=1, sticky='nsew', padx=(12, 0))
        self.update_excel_button_state()

    def update_excel_button_state(self):
        if not self.generate_excel_btn:
            return
        if self.excel_path and not self.generating_var.get():
            self.generate_excel_btn.config(state='normal')
        else:
            self.generate_excel_btn.config(state='disabled')
        
    # File selection methods
    def select_template(self):
        file_path = filedialog.askopenfilename(
            title="Select Word Template",
            filetypes=[("Word Documents", "*.docx")]
        )
        if file_path:
            self.template_path = file_path
            self.template_name_var.set(Path(file_path).name)
            self.show_toast(f"Template loaded: {Path(file_path).name}")
            self.update_status("Template ready")
            
    def select_excel(self):
        file_path = filedialog.askopenfilename(
            title="Select Excel File",
            filetypes=[("Excel Files", "*.xlsx *.xls")]
        )
        if file_path:
            self.excel_path = file_path
            self.excel_name_var.set(Path(file_path).name)
            self.show_toast(f"Excel loaded: {Path(file_path).name}")
            self.update_status("Excel ready")
            self.update_excel_button_state()
            
    def select_folder(self):
        folder_path = filedialog.askdirectory(title="Select Output Folder")
        if folder_path:
            self.output_folder = folder_path
            self.output_folder_var.set(Path(folder_path).name)
            self.show_toast(f"Output folder set: {Path(folder_path).name}")
            self.update_status("Output folder ready")
            
    def scan_placeholders(self):
        if not self.template_path or not self.excel_path:
            self.show_toast("Select both files first")
            return
            
        try:
            self.update_status("Scanning placeholders...")
            
            # Scan template for placeholders
            doc = Document(self.template_path)
            placeholders = set()
            regex = re.compile(r"\{\{\s*([A-Za-z0-9_]+)\s*\}\}")
            
            for p in doc.paragraphs:
                placeholders.update(regex.findall(p.text))
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        placeholders.update(regex.findall(cell.text))
                        
            # Get Excel columns
            df = pd.read_excel(self.excel_path, nrows=0)
            columns = list(df.columns)
            
            self.placeholders = sorted(placeholders)
            self.columns = columns
            self.mapping = {}
            
            self.update_mapping_display()
            self.update_status(f"Found {len(self.placeholders)} fields")
            self.show_toast("Scan complete!")
            
        except Exception as e:
            self.show_toast("Scan failed")
            print(f"Scan error: {e}")
            
    def update_mapping_display(self):
        # Clear existing widgets
        for widget in self.scrollable_frame.winfo_children():
            widget.destroy()
            
        # Update field count
        self.field_count_label.config(text=f"{len(self.placeholders)} fields → {len(self.columns)} columns")
        
        if not self.placeholders:
            self.show_empty_state()
            return
            
        # Create mapping widgets
        for i, placeholder in enumerate(self.placeholders):
            self.create_mapping_row(i, placeholder)
            
    def create_mapping_row(self, index, placeholder):
        row_frame = tk.Frame(self.scrollable_frame, bg='white', relief='solid', bd=1)
        row_frame.pack(fill='x', padx=20, pady=5)
        
        # Index number (Caption: 12px)
        index_label = tk.Label(row_frame, text=str(index + 1), 
                              font=('Segoe UI', 12, 'bold'), 
                              bg='#3b82f6', fg='white', width=4)
        index_label.pack(side='left', padx=(12, 12), pady=12)
        
        # Placeholder name (Body: 14px)
        placeholder_label = tk.Label(row_frame, text=f"{{ {placeholder} }}", 
                                   font=('Consolas', 14, 'bold'), 
                                   bg='#f3f4f6', fg='#374151', width=22)
        placeholder_label.pack(side='left', padx=(0, 12), pady=12)
        
        # Dropdown for column selection (Body: 14px)
        column_var = tk.StringVar()
        column_combo = ttk.Combobox(row_frame, textvariable=column_var, 
                                   values=["— Select Excel Column —"] + self.columns,
                                   state="readonly", font=('Segoe UI', 14), width=32)
        
        # Set initial value based on mapping
        if placeholder in self.mapping:
            column_combo.set(self.mapping[placeholder])
        else:
            column_combo.set("— Select Excel Column —")
            
        column_combo.pack(side='left', fill='x', expand=True, padx=(0, 12), pady=12)
        
        # Status indicator
        status_label = tk.Label(row_frame, text="○", font=('Segoe UI', 16), 
                               fg='#d1d5db', bg='white')
        
        # Set initial status based on mapping
        if placeholder in self.mapping:
            status_label.config(text="●", fg='#10b981')
        else:
            status_label.config(text="○", fg='#d1d5db')
            
        status_label.pack(side='right', padx=12, pady=12)
        
        # Bind selection event
        def on_select(event, ph=placeholder, var=column_var, status=status_label):
            selected = var.get()
            if selected and selected != "— Select Excel Column —":
                self.mapping[ph] = selected
                status.config(text="●", fg='#10b981')
            else:
                if ph in self.mapping:
                    del self.mapping[ph]
                status.config(text="○", fg='#d1d5db')
                
        column_combo.bind('<<ComboboxSelected>>', on_select)
        
    def auto_map(self):
        if not self.placeholders or not self.columns:
            self.show_toast("Scan placeholders first")
            return
            
        self.update_status("Auto-mapping fields...")
        
        # Create column mapping with multiple variations
        col_map = {}
        for col in self.columns:
            variations = [
                col.lower(),
                col.lower().replace(' ', '_'),
                col.lower().replace('_', ' '),
                col.lower().replace('-', '_'),
                col.lower().replace('.', '_'),
                ''.join(col.lower().split()),  # Remove all spaces
            ]
            for var in variations:
                col_map[var] = col
        
        mapped_count = 0
        
        for placeholder in self.placeholders:
            if placeholder in self.mapping:
                continue  # Skip already mapped
                
            ph_lower = placeholder.lower()
            
            # Try multiple matching strategies
            match_found = False
            
            # 1. Exact match
            if ph_lower in col_map:
                self.mapping[placeholder] = col_map[ph_lower]
                mapped_count += 1
                match_found = True
            
            # 2. Try variations of placeholder
            if not match_found:
                ph_variations = [
                    ph_lower.replace('_', ' '),
                    ph_lower.replace(' ', '_'),
                    ph_lower.replace('-', '_'),
                    ''.join(ph_lower.split()),
                ]
                
                for ph_var in ph_variations:
                    if ph_var in col_map:
                        self.mapping[placeholder] = col_map[ph_var]
                        mapped_count += 1
                        match_found = True
                        break
            
            # 3. Partial matching (contains)
            if not match_found:
                for col in self.columns:
                    col_lower = col.lower()
                    if (ph_lower in col_lower or col_lower in ph_lower) and len(ph_lower) > 2:
                        self.mapping[placeholder] = col
                        mapped_count += 1
                        break
                
        self.update_mapping_display()
        self.show_toast(f"Auto-mapped {mapped_count} fields")
        self.update_status(f"Mapped {mapped_count}/{len(self.placeholders)} fields")
        
    def clear_mapping(self):
        self.mapping = {}
        self.update_mapping_display()
        self.show_toast("All mappings cleared")
        self.update_status("Mappings cleared")
        
    def preview_data(self):
        if not self.excel_path:
            self.show_toast("Select Excel file first")
            return
            
        try:
            df = pd.read_excel(self.excel_path, dtype=str).fillna("")
            if df.empty:
                messagebox.showinfo("Preview", "Excel file is empty")
                return
                
            row = df.iloc[0].to_dict()
            lines = []
            
            for placeholder, column in self.mapping.items():
                if column:
                    value = row.get(column, "(empty)")
                    lines.append(f"{placeholder} = {value}")
                    
            if lines:
                preview_text = "Preview First Row:\n\n" + "\n".join(lines)
            else:
                preview_text = "No mappings configured yet"
                
            messagebox.showinfo("Data Preview", preview_text)
            
        except Exception as e:
            messagebox.showerror("Preview Error", f"Failed to preview data: {str(e)}")
            
    def generate_documents(self):
        if not self.mapping:
            self.show_toast("Please map at least one field before generating")
            return
            
        if self.generating_var.get():
            return
            
        self.generating_var.set(True)
        self.generate_btn.config(text="🔄 Generating...", state='disabled')
        self.update_excel_button_state()
        
        def run_generation():
            try:
                self.update_status("Generating documents...")
                
                df = pd.read_excel(self.excel_path, dtype=str).fillna("")
                tpl = DocxTemplate(self.template_path)
                
                docx_dir = Path(self.output_folder) / "DocGen-Engine_DOCX"
                pdf_dir = Path(self.output_folder) / "DocGen-Engine_PDF"
                docx_dir.mkdir(exist_ok=True)
                pdf_dir.mkdir(exist_ok=True)
                
                name_counts = {}
                generated_names = []
                for idx, row in df.iterrows():
                    if self._is_row_empty(row):
                        generated_names.append("")
                        continue
                    context = {ph: str(row.get(col, "")) for ph, col in self.mapping.items()}
                    tpl.render(context)
                    
                    name = self._build_filename(row, idx, name_counts)
                    generated_names.append(name)
                    tpl.save(str(docx_dir / f"{name}.docx"))
                    
                self.update_status("Converting to PDF...")
                
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                try:
                    word.Options.DoNotCompressImagesInfile = True
                except Exception:
                    pass  # Older Word versions might not expose this option
                
                for docx_path in docx_dir.glob("*.docx"):
                    pdf_path = str(pdf_dir / (docx_path.stem + ".pdf"))
                    doc = word.Documents.Open(str(docx_path.resolve()))
                    try:
                        try:
                            doc.SaveAs2(
                                str(docx_path.resolve()),
                                FileFormat=16,  # wdFormatXMLDocument
                                EmbedTrueTypeFonts=True,
                                AddToRecentFiles=False,
                                DoNotCompressImagesInfile=True,
                            )
                        except Exception:
                            pass
                        doc.ExportAsFixedFormat(
                            OutputFileName=pdf_path,
                            ExportFormat=17,  # wdExportFormatPDF
                            OpenAfterExport=False,
                            OptimizeFor=0,  # wdExportOptimizeForPrint (highest quality)
                            Range=0,  # wdExportAllDocument
                            Item=0,  # wdExportDocumentContent
                            IncludeDocProps=True,
                            KeepIRM=True,
                            CreateBookmarks=1,  # wdExportCreateHeadingBookmarks
                            DocStructureTags=True,
                            BitmapMissingFonts=True,
                            UseISO19005_1=False,
                        )
                    except Exception:
                        # Fallback to SaveAs if ExportAsFixedFormat is unavailable
                        doc.SaveAs(pdf_path, FileFormat=17)
                    finally:
                        doc.Close()
                    
                word.Quit()
                
                result_df = df.copy()
                self.root.after(0, lambda: self.generation_complete(docx_dir, pdf_dir, result_df, generated_names))
                
            except Exception as e:
                traceback.print_exc()
                error_msg = str(e)
                self.root.after(0, lambda msg=error_msg: self.generation_error(msg))
                
        threading.Thread(target=run_generation, daemon=True).start()
        
    def generation_complete(self, docx_dir, pdf_dir, dataframe=None, generated_names=None):
        self.generating_var.set(False)
        self.generate_btn.config(text="🚀 Generate DOCX & PDF Files", state='normal')
        self.update_status("All done!")
        self.show_toast("Generation Complete!")
        self.update_excel_button_state()
        messagebox.showinfo(
            "Success!",
            "Documents generated successfully!\n\nDOCX: {}\nPDF: {}\n\nUse 'Generate Excel Report' to export the summary.".format(docx_dir, pdf_dir)
        )
        
    def generation_error(self, error_msg):
        self.generating_var.set(False)
        self.generate_btn.config(text="🚀 Generate DOCX & PDF Files", state='normal')
        self.update_status("Error occurred")
        messagebox.showerror("Generation Error", f"Error: {error_msg}")
        self.update_excel_button_state()
        
    def generate_excel_report(self):
        if not self.excel_path:
            messagebox.showinfo("Export", "Select an Excel file first.")
            return
        try:
            df = pd.read_excel(self.excel_path, dtype=str).fillna("")
            if df.empty:
                messagebox.showinfo("Export", "Excel file is empty. Nothing to export.")
                return
            generated_names = self._compute_filenames(df)
            report_df = df.copy()
            report_df["Generated File"] = [f"{name}.pdf" if name else "" for name in generated_names]
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            report_path = Path(self.output_folder) / f"DocGen-Engine_Report_{timestamp}.xlsx"
            report_df.to_excel(report_path, index=False)
            self.show_toast(f"Report exported: {report_path}")
            messagebox.showinfo("Export", f"Summary saved to:\n{report_path}")
        except Exception as e:
            traceback.print_exc()
            messagebox.showerror("Export", f"Failed to export summary: {e}")
    
    def _compute_filenames(self, df):
        name_counts = {}
        filenames = []
        for idx, row in df.iterrows():
            if self._is_row_empty(row):
                filenames.append("")
                continue
            filenames.append(self._build_filename(row, idx, name_counts))
        return filenames
    
    def _build_filename(self, row, idx, name_counts):
        name = str(row.iloc[0]).strip() if len(row) > 0 else f"Document_{idx+1}"
        name = re.sub(r"[^\w\-_.]", "_", name) or "Document"
        
        if self.include_mobile_var.get():
            mobile = ""
            for col in row.index:
                col_lower = str(col).lower()
                if any(x in col_lower for x in ["mobile", "phone", "contact"]):
                    m = re.sub(r"\D", "", str(row[col]))
                    if len(m) >= 10:
                        mobile = m[-10:]
                        break
            if mobile:
                name += "_" + mobile
        
        base_name = name
        count = name_counts.get(base_name, 0) + 1
        name_counts[base_name] = count
        if count > 1:
            name = f"{base_name}_{count}"
        return name

    def _is_row_empty(self, row):
        for value in row:
            if str(value).strip():
                return False
        return True
        
    def show_toast(self, message):
        # Simple status update (could be enhanced with actual toast notifications)
        print(f"Toast: {message}")
        
    def update_status(self, status):
        self.status_var.set(status)
        
    def run(self):
        self.root.mainloop()

if __name__ == "__main__":
    app = SaralWorksApp()
    app.run()